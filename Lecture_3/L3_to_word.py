import locale
import time

import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 设置对齐方式
from docx.oxml.ns import qn  # 设置字体
from docx.shared import Cm, Inches, Pt, RGBColor  # 用来设置字体大小、缩进; 设置字体的颜色

locale.setlocale(locale.LC_CTYPE, "chinese")
date = time.strftime("%Y年%m月%d日")


# 使用Pandas读取excel文件
data = pd.read_excel("活动数据.xlsx")

# 封装成函数，批量生产文档
def denerate_docx(info):

    gifts = {"一等奖": "一台Kindle", "二等奖": "小米手环", "三等奖": "一沓数学作业纸"}

    prize = ""
    if info[4] == 200:
        prize = "一等奖"
    elif info[4] >= 190:
        prize = "二等奖"
    elif info[4] >= 180:
        prize = "三等奖"
    else:
        return

    doc = docx.Document()  # 创建文档对象

    heading = doc.add_paragraph("")  # 创建段落对象，当作标题
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置为居中对齐

    title = "公示"
    run = heading.add_run(title)  # 添加标题文字

    run.font.name = "黑体"  # 设置字体
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")  # 设置为字体，和上边的保持一致
    run.font.size = Pt(20)  # 设置文字的大小为20磅
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色

    # 设置正文格式
    doc.styles["Normal"].font.size = Pt(14)
    doc.styles["Normal"].font.name = "仿宋"  # 设置当文字是西文时的字体
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")  # 设置当文字是中文时的字体

    # 添加正文信息
    doc.add_paragraph("")  # 增加空段落
    text = f"根据海量专家的精密计算，{info[0]}同学在本次活动中成功打卡{info[4]}天，GPA达到了{info[5]:0.4f}，获得{prize}，奖品是{gifts[prize]}。特此公示。"
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Cm(1)  # 设置左缩进 1 英寸

    text2 = "公示期自即日始5个工作日，凡对上述同学获奖有意见者，请及时以书面或口头形式向XX大学学生学业发展中心反映。"
    paragraph = doc.add_paragraph(text2)
    paragraph.paragraph_format.first_line_indent = Cm(1)  # 设置左缩进 1 英寸

    doc.add_paragraph("")  # 增加空段落

    paragraph = doc.add_paragraph("XX大学学生学业发展中心")  # 落款
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐

    paragraph = doc.add_paragraph(date)  # 日期
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 右对齐

    # 保存文档（保存时请确认该文档是关闭的，否则会出权限错误）
    doc.save(f"output/公示_{prize}_{info[0]}.docx")


for index, row in data.iterrows():
    denerate_docx(row)

print("文档生成成功！")
